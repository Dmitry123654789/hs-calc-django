from datetime import date
from decimal import Decimal, InvalidOperation
import io
import os
from typing import Iterable, Optional, Sequence

from django.contrib.staticfiles import finders
from docx import Document
from docx.document import Document as DocumentType
from docx.enum.table import WD_ROW_HEIGHT_RULE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from PIL import Image


def asset_path(name: str) -> str:
    path = finders.find(f"img/calculate/{name}")
    if isinstance(path, list):
        return str(path[0]) if path else ""

    return str(path) if path else ""


GLUKHAR_IMAGE = asset_path("schemes/glukhar_default.jpg")
COMPANY_LOGO = asset_path("logo.png")

SCHEME_IMAGE_MAP = {
    "scheme_A": asset_path("schemes/scheme_A.png"),
    "scheme_C": asset_path("schemes/scheme_C.png"),
    "scheme_G": asset_path("schemes/scheme_G.png"),
    "scheme_E": asset_path("schemes/scheme_E.png"),
    "scheme_L": asset_path("schemes/scheme_L.png"),
}

# Сторона ручки по названию типа фурнитуры (Hardware.name)
HARDWARE_SIDE_MAP = {
    "Standard": "ручка односторонняя",
    "Standard+": "ручка двухсторонняя",
}

GLASS_TYPE_TEXT = "40мм закалённый"

FONT_NAME = "Calibri"
NORMAL_FONT_SIZE = 14
HEADING_FONT_SIZE = 18

PAGE_WIDTH_CM = 21.0
PAGE_HEIGHT_CM = 29.7


def _to_decimal(value) -> Decimal:
    if isinstance(value, Decimal):
        return value

    if value is None:
        return Decimal("0")

    try:
        return Decimal(str(value))

    except (InvalidOperation, ValueError, TypeError):
        return Decimal("0")


def _fmt_money(value) -> str:
    """1999999.5 -> '1 999 999,50 ₽'"""
    value = _to_decimal(value).quantize(Decimal("0.01"))
    int_part = str(abs(value.to_integral_value(rounding="ROUND_FLOOR")))
    frac = f"{abs(value):.2f}".split(".")[1]
    grouped = []
    while len(int_part) > 3:
        grouped.insert(0, int_part[-3:])
        int_part = int_part[:-3]

    grouped.insert(0, int_part)
    formatted = " ".join(grouped)
    prefix = "-" if value < 0 else ""
    return f"{prefix}{formatted},{frac} ₽"


def _fmt_area(value) -> str:
    value = _to_decimal(value).quantize(Decimal("0.01"))
    return f"{value}".replace(".", ",")


def get_product_price(product) -> Decimal:
    """
    Достаёт итоговую (с коэффициентом) стоимость партии изделия из уже
    посчитанного product.calculation_details. Portal и Glukhar отдают
    одинаковую по форме структуру и всегда кладут туда ключ
    "total_with_ratio", поэтому обработка — прямое обращение по ключу,
    без перебора возможных вариантов названия.
    """
    details = getattr(product, "calculation_details", None) or {}
    if not isinstance(details, dict):
        return Decimal("0")

    return _to_decimal(details.get("total_with_ratio"))


def get_portal_image_path(portal) -> str:
    scheme_name = getattr(getattr(portal, "scheme", None), "name", None)
    return SCHEME_IMAGE_MAP.get(scheme_name) or ""


def get_hardware_side_text(portal) -> str:
    hardware_name = getattr(getattr(portal, "hardware_type", None), "name", None)
    return HARDWARE_SIDE_MAP.get(hardware_name, "ручка односторонняя")


def get_color_display(product) -> str:
    color = getattr(product, "color_type", None)
    return getattr(color, "name", "") or "-"


def get_scheme_letter(scheme) -> str:
    """
    Scheme.name в БД всегда имеет вид "scheme_<буква>" (scheme_A, scheme_C,
    scheme_G, scheme_E, scheme_L, ...) — буква для заголовка "Схема А." это
    то, что стоит после последнего "_" в имени.
    """
    return scheme.name.rsplit("_", 1)[-1]


# Низкоуровневые помощники для python-docx
def _set_font(run, size=NORMAL_FONT_SIZE, bold=False, color=None):
    run.font.name = FONT_NAME
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)

    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = rpr.makeelement(qn("w:rFonts"), {})
        rpr.append(rfonts)

    rfonts.set(qn("w:eastAsia"), FONT_NAME)


def _configure_normal_style(doc: DocumentType) -> None:
    """Шрифт Calibri/14pt и нулевые интервалы абзаца как базовый стиль
    документа — подстраховка для параграфов/ячеек, где не выставляется
    прямое форматирование запуска (run)."""
    normal = doc.styles["Normal"]
    normal.font.name = FONT_NAME
    normal.font.size = Pt(NORMAL_FONT_SIZE)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(0)
    rpr = normal.element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = rpr.makeelement(qn("w:rFonts"), {})
        rpr.append(rfonts)

    rfonts.set(qn("w:eastAsia"), FONT_NAME)


def _zero_spacing(paragraph) -> None:
    """Интервал перед/после абзаца = 0pt — применяется к каждому создаваемому
    абзацу, чтобы гарантированно выполнялось «везде», а не только там, где
    сработал стиль Normal (прямое форматирование имеет приоритет)."""
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)


def _set_paragraph_edge_indent(paragraph, left_cm: float, right_cm: float) -> None:
    paragraph.paragraph_format.left_indent = Cm(left_cm)
    paragraph.paragraph_format.right_indent = Cm(right_cm)


def _set_table_left_indent(table, left_cm: float) -> None:
    """Абсолютный отступ таблицы от левого края (w:tblInd). Требует, чтобы
    поля секции (left/right margin) были равны 0 — иначе отступ будет
    считаться от границы текстовой области, а не от края листа."""
    tbl_pr = table._tbl.tblPr
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)

    tbl_ind.set(qn("w:w"), str(Cm(left_cm).twips))
    tbl_ind.set(qn("w:type"), "dxa")


def _zero_table_cell_margins(table) -> None:
    """
    Обнуляет внутренние поля ячеек по умолчанию для всей таблицы.
    Без этого Word/LibreOffice подставляют свои дефолтные поля ячейки
    (~0.19-0.25 см с каждой стороны), и фактическое содержимое ячейки
    (в частности — картинка ровно по ширине ячейки) визуально сжимается
    внутрь на эту величину, из-за чего реальный отступ от края листа
    отличается от заданного.
    """
    tbl_pr = table._tbl.tblPr
    cell_mar = tbl_pr.find(qn("w:tblCellMar"))
    if cell_mar is None:
        cell_mar = OxmlElement("w:tblCellMar")
        tbl_pr.append(cell_mar)

    for side in ("top", "left", "bottom", "right"):
        el = cell_mar.find(qn(f"w:{side}"))
        if el is None:
            el = OxmlElement(f"w:{side}")
            cell_mar.append(el)

        el.set(qn("w:w"), "0")
        el.set(qn("w:type"), "dxa")


def _set_table_width_dxa(table, width_dxa: int) -> None:
    """Явно фиксирует общую ширину таблицы (w:tblW) в twips, вместо
    оставленного по умолчанию type="auto" — иначе рендерер может сам
    пересчитывать итоговую ширину таблицы, а не брать её строго из суммы
    ширин колонок."""
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)

    tbl_w.set(qn("w:w"), str(width_dxa))
    tbl_w.set(qn("w:type"), "dxa")


def _set_table_edge_indent(
    table,
    left_cm: float,
    right_cm: float,
    col_widths_cm: Sequence[float],
    *,
    zero_cell_margins: bool = False,
) -> None:
    """
    Позиционирует таблицу строго по отступам от левого/правого края листа:
    задаёт w:tblInd = left_cm, а ширину колонок (сумма = PAGE_WIDTH_CM -
    left_cm - right_cm) выставляет и на уровне колонок, и на уровне каждой
    ячейки (без дублирования на ячейках python-docx/Word не всегда
    применяет ширину колонки), плюс явную общую ширину таблицы (иначе
    рендерер может пересчитать её сам вместо суммы колонок).
    zero_cell_margins=True дополнительно обнуляет внутренние поля ячеек —
    нужно там, где содержимое (например картинка) обязано занимать ровно
    заявленную ширину ячейки без «утапливания» дефолтными полями Word
    (~0.2 см с каждой стороны). Для обычных текстовых таблиц оставляем
    поля Word по умолчанию — иначе текст визуально прилипает к границам
    ячеек.
    Таблица не должна иметь w:jc (table.alignment) — это конфликтует с
    абсолютным отступом tblInd.
    """
    table.autofit = False
    _set_table_left_indent(table, left_cm)
    if zero_cell_margins:
        _zero_table_cell_margins(table)

    total_dxa = 0
    for i, w in enumerate(col_widths_cm):
        table.columns[i].width = Cm(w)
        total_dxa += Cm(w).twips

    for row in table.rows:
        for cell, w in zip(row.cells, col_widths_cm):
            cell.width = Cm(w)

    _set_table_width_dxa(table, total_dxa)


def _set_cell_borders(
    cell,
    *,
    top: Optional[str] = None,
    bottom: Optional[str] = None,
    left: Optional[str] = None,
    right: Optional[str] = None,
) -> None:
    """
    Управляет границами конкретной ячейки по каждой стороне отдельно.
    Значение стороны: None — не трогать (наследуется от стиля таблицы),
    "none" — скрыть эту границу у этой ячейки, "single" — одинарная
    чёрная линия 0.5pt (как в стиле Table Grid).
    Прямое форматирование на уровне ячейки имеет приоритет над стилем
    таблицы, поэтому этим можно сделать один столбец полностью без
    рамок, а соседний — полностью с рамками, даже в одной таблице.
    """
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.find(qn("w:tcBorders"))
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)

    for side, spec in (
        ("top", top),
        ("bottom", bottom),
        ("left", left),
        ("right", right),
    ):
        if spec is None:
            continue

        el = borders.find(qn(f"w:{side}"))
        if el is None:
            el = OxmlElement(f"w:{side}")
            borders.append(el)

        if spec == "none":
            el.set(qn("w:val"), "nil")
            el.set(qn("w:sz"), "0")
            el.set(qn("w:space"), "0")
            el.set(qn("w:color"), "auto")
        elif spec == "single":
            el.set(qn("w:val"), "single")
            el.set(qn("w:sz"), "4")  # 4 = 0.5pt, восьмые доли пункта
            el.set(qn("w:space"), "0")
            el.set(qn("w:color"), "000000")


def _add_heading(
    doc: DocumentType,
    text: str,
    *,
    bold: bool = True,
    size: int = NORMAL_FONT_SIZE,
    left_cm: float = 1.5,
    right_cm: float = 1.5,
) -> None:
    p = doc.add_paragraph()
    _zero_spacing(p)
    _set_paragraph_edge_indent(p, left_cm - 0.2, right_cm)
    run = p.add_run(text)
    _set_font(run, size=size, bold=bold)


def _add_kv_table(
    doc: DocumentType,
    rows: Iterable[tuple[str, str]],
    *,
    left_cm: float = 1.5,
    right_cm: float = 1.5,
) -> None:
    """Двухколоночная таблица 'подпись | значение' (Материал/Фурнитура/Цвет/...)
    — «таблица с характеристиками товара»."""
    rows = list(rows)
    table = doc.add_table(rows=len(rows), cols=2)
    table.style = "Table Grid"

    total_width = PAGE_WIDTH_CM - left_cm - right_cm
    label_w = total_width * (4.5 / 16)
    value_w = total_width - label_w
    _set_table_edge_indent(table, left_cm, right_cm, [label_w, value_w])

    for i, (label, value) in enumerate(rows):
        row = table.rows[i]

        p0 = row.cells[0].paragraphs[0]
        _zero_spacing(p0)
        r0 = p0.add_run(label)
        _set_font(r0, bold=True)

        p1 = row.cells[1].paragraphs[0]
        _zero_spacing(p1)
        r1 = p1.add_run(str(value))
        _set_font(r1)

    spacer = doc.add_paragraph()
    _zero_spacing(spacer)


def _safe_image_stream(image_path: str) -> tuple[io.BytesIO, int, int]:
    """
    python-docx's own JPEG-header парсер иногда не распознаёт progressive
    JPEG (частый случай для картинок из static/CMS-выгрузок) и падает с
    UnrecognizedImageError, хотя файл абсолютно валиден. Прогоняем через
    Pillow и пересохраняем в надёжном baseline-формате перед вставкой.
    Заодно отдаём исходные пиксельные размеры — нужны, чтобы вписать
    изображение в ячейку с сохранением пропорций.
    """
    with Image.open(image_path) as im:
        width_px, height_px = im.size
        im = im.convert("RGB") if im.mode not in ("RGB", "RGBA") else im
        buf = io.BytesIO()
        fmt = "PNG" if im.mode == "RGBA" else "JPEG"
        im.save(buf, format=fmt)
        buf.seek(0)
        return buf, width_px, height_px


def _fit_box(
    img_w_px: int,
    img_h_px: int,
    box_w_cm: float,
    box_h_cm: float,
) -> tuple[float, float]:
    """Пропорциональное вписывание картинки в прямоугольник box_w x box_h
    (касание только двух стенок, без обрезки и искажения — classic
    "contain" fit)."""
    img_ratio = img_w_px / img_h_px
    box_ratio = box_w_cm / box_h_cm
    if img_ratio >= box_ratio:
        return box_w_cm, box_w_cm / img_ratio

    return box_h_cm * img_ratio, box_h_cm


PRODUCT_IMAGE_BOX_W_CM = 7.77
PRODUCT_IMAGE_BOX_H_CM = 6.04

PRODUCT_IMAGE_FIT_MARGIN_CM = 0.3
PRODUCT_IMAGE_FIT_W_CM = PRODUCT_IMAGE_BOX_W_CM - PRODUCT_IMAGE_FIT_MARGIN_CM
PRODUCT_IMAGE_FIT_H_CM = PRODUCT_IMAGE_BOX_H_CM - PRODUCT_IMAGE_FIT_MARGIN_CM


def _add_product_block(
    doc: DocumentType,
    image_path: str,
    spec_rows: Iterable[tuple[str, str]],
    *,
    left_cm: float = 1.5,
    right_cm: float = 1.5,
) -> None:
    """
    Таблица 'картинка | подпись | значение' — картинка изделия слева
    (объединена по вертикали на все строки), справа характеристики.
    Это «таблица с описанием товара».
    """
    spec_rows = list(spec_rows)
    n_rows = len(spec_rows)

    table = doc.add_table(rows=n_rows, cols=3)
    table.style = "Table Grid"

    total_width = PAGE_WIDTH_CM - left_cm - right_cm
    img_col_w = PRODUCT_IMAGE_BOX_W_CM
    remaining_w = total_width - img_col_w
    label_col_w = remaining_w * 0.6
    value_col_w = remaining_w - label_col_w
    _set_table_edge_indent(
        table,
        left_cm,
        right_cm,
        [img_col_w, label_col_w, value_col_w],
    )

    # Точная высота строк, чтобы объединённая ячейка с фото была ровно 6.04 см
    row_height = Cm(PRODUCT_IMAGE_BOX_H_CM / n_rows)
    for row in table.rows:
        row.height = row_height
        row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY

    # Объединяем картиночную колонку по вертикали
    img_cell = table.cell(0, 0)
    if n_rows > 1:
        img_cell = img_cell.merge(table.cell(n_rows - 1, 0))

    img_cell.width = Cm(img_col_w)
    img_cell.vertical_alignment = 1  # center
    img_para = img_cell.paragraphs[0]
    _zero_spacing(img_para)
    img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = img_para.add_run()
    try:
        stream, px_w, px_h = _safe_image_stream(image_path)
        fit_w, fit_h = _fit_box(
            px_w,
            px_h,
            PRODUCT_IMAGE_FIT_W_CM,
            PRODUCT_IMAGE_FIT_H_CM,
        )
        run.add_picture(stream, width=Cm(fit_w), height=Cm(fit_h))
    except Exception:
        # если файл картинки не найден/битый — не роняем генерацию документа
        placeholder = img_para.add_run("[изображение недоступно]")
        _set_font(placeholder, size=NORMAL_FONT_SIZE)

    for i, (label, value) in enumerate(spec_rows):
        label_cell = table.cell(i, 1)
        value_cell = table.cell(i, 2)

        lp = label_cell.paragraphs[0]
        _zero_spacing(lp)
        lr = lp.add_run(label)
        _set_font(lr)

        vp = value_cell.paragraphs[0]
        _zero_spacing(vp)
        vr = vp.add_run(str(value))
        _set_font(vr, bold=True)

    spacer = doc.add_paragraph()
    _zero_spacing(spacer)


# Сборка блоков "Портал" и "Глухарь"
def _add_portal_section(doc: DocumentType, portal) -> Decimal:
    scheme = portal.scheme
    letter = get_scheme_letter(scheme)
    _add_heading(doc, f"Схема {letter}. HS-портал", size=HEADING_FONT_SIZE)

    width = portal.width
    height = portal.height
    amount = portal.amount
    area_total = (
        _to_decimal(width) * _to_decimal(height) / Decimal("1000000")
    ) * _to_decimal(amount)
    price = get_product_price(portal)

    spec_rows = [
        ("Высота проема, мм:", str(height)),
        ("Ширина проема, мм:", str(width)),
        ("Количество изделий, шт", str(amount)),
        ("Площадь изделий, м\u00b2", _fmt_area(area_total)),
        ("Стоимость изделий, рублей", _fmt_money(price)),
    ]
    _add_product_block(doc, get_portal_image_path(portal), spec_rows)

    wood_name = getattr(portal.wood_type, "name", "-")
    _add_kv_table(
        doc,
        [
            ("Материал", f"{wood_name}, клееный брус 78 мм"),
            ("Фурнитура", f"HS-Portal, {get_hardware_side_text(portal)}"),
            ("Цвет", get_color_display(portal)),
            ("Тип стеклопакета", GLASS_TYPE_TEXT),
        ],
    )
    return price


def _add_glukhar_section(doc: DocumentType, glukhar, index: int) -> Decimal:
    _add_heading(doc, "Глухое окно", size=HEADING_FONT_SIZE)

    width = glukhar.width
    height = glukhar.height
    amount = glukhar.amount
    area_total = (
        _to_decimal(width) * _to_decimal(height) / Decimal("1000000")
    ) * _to_decimal(amount)
    price = get_product_price(glukhar)

    spec_rows = [
        ("Высота проема, мм:", str(height)),
        ("Ширина проема, мм:", str(width)),
        ("Количество изделий, шт", str(amount)),
        ("Площадь изделий, м\u00b2", _fmt_area(area_total)),
        ("Стоимость изделий, рублей", _fmt_money(price)),
    ]
    _add_product_block(doc, GLUKHAR_IMAGE, spec_rows)

    wood_name = getattr(glukhar.wood_type, "name", "-")
    _add_kv_table(
        doc,
        [
            ("Материал", f"{wood_name}, клееный брус"),
            ("Цвет", get_color_display(glukhar)),
            ("Тип стеклопакета", GLASS_TYPE_TEXT),
        ],
    )
    return price


# Шапка и итоговый блок
HEADER_TEXT_LEFT_CM = 2.5
HEADER_TEXT_RIGHT_CM = 1.5
LOGO_CELL_CM = 2.75
LOGO_TABLE_RIGHT_CM = 1.5

DATES_LEFT_CM = 2.5
DATES_RIGHT_CM = 1.5
DATES_LEFT_COL_CM = 3.5


def _add_document_header(
    doc: DocumentType,
    company_name: str,
    inn: str,
    ogrn: str,
    proposal_date: str,
    shipment_term: str,
) -> None:
    # Блок реквизитов (слева) + блок с логотипом (справа) — одна строка,
    # позиционированная так, чтобы текстовый блок начинался в 2.5 см от
    # левого края листа, а таблица с логотипом заканчивалась в 1.5 см от
    # правого края листа (оба требования выполняются одновременно, т.к.
    # это один и тот же ряд с общими левым/правым отступами).
    table = doc.add_table(rows=1, cols=2)
    total_width = PAGE_WIDTH_CM - HEADER_TEXT_LEFT_CM - LOGO_TABLE_RIGHT_CM
    text_col_w = total_width - LOGO_CELL_CM
    _set_table_edge_indent(
        table,
        HEADER_TEXT_LEFT_CM,
        LOGO_TABLE_RIGHT_CM,
        [text_col_w - 1.5, LOGO_CELL_CM + 1.5],
    )

    left_cell = table.cell(0, 0)
    p = left_cell.paragraphs[0]
    _zero_spacing(p)
    r = p.add_run("КОММЕРЧЕСКОЕ ПРЕДЛОЖЕНИЕ")
    _set_font(r, bold=True)

    p2 = left_cell.add_paragraph()
    _zero_spacing(p2)
    r2 = p2.add_run(company_name)
    _set_font(r2)

    p3 = left_cell.add_paragraph()
    _zero_spacing(p3)
    r3 = p3.add_run(f"ИНН {inn}")
    _set_font(r3)

    p4 = left_cell.add_paragraph()
    _zero_spacing(p4)
    r4 = p4.add_run(f"ОГРН {ogrn}")
    _set_font(r4)

    right_cell = table.cell(0, 1)
    right_cell.width = Cm(LOGO_CELL_CM)
    rp = right_cell.paragraphs[0]
    _zero_spacing(rp)
    rp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rrun = rp.add_run()
    if COMPANY_LOGO and os.path.exists(COMPANY_LOGO):
        try:
            rrun.add_picture(COMPANY_LOGO, width=Cm(LOGO_CELL_CM))
        except Exception:
            pass

    spacer = doc.add_paragraph()
    _zero_spacing(spacer)

    # Таблица с датами («Дата КП:» / «Срок отгрузки:»)
    dates_total_width = PAGE_WIDTH_CM - DATES_LEFT_CM - DATES_RIGHT_CM
    dates_right_col_cm = dates_total_width - DATES_LEFT_COL_CM

    meta_table = doc.add_table(rows=2, cols=2)
    meta_table.style = "Table Grid"
    _set_table_edge_indent(
        meta_table,
        DATES_LEFT_CM,
        DATES_RIGHT_CM,
        [DATES_LEFT_COL_CM, dates_right_col_cm],
    )

    meta_rows = [("Дата КП:", proposal_date), ("Срок отгрузки:", shipment_term)]
    for i, (label, value) in enumerate(meta_rows):
        c0, c1 = meta_table.rows[i].cells

        # Левый столбец — совсем без границ, правый — со всеми границами.
        _set_cell_borders(c0, top="none", bottom="none", left="none", right="single")
        _set_cell_borders(
            c1,
            top="single",
            bottom="single",
            left="single",
            right="single",
        )

        p0 = c0.paragraphs[0]
        _zero_spacing(p0)
        p0.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        r0 = p0.add_run(label)
        _set_font(r0)

        p1 = c1.paragraphs[0]
        _zero_spacing(p1)
        p1.alignment = WD_ALIGN_PARAGRAPH.LEFT
        r1 = p1.add_run(value)
        _set_font(r1)

    spacer2 = doc.add_paragraph()
    _zero_spacing(spacer2)


COST_LEFT_CM = 1.5
COST_RIGHT_CM = 1.5
PAYMENT_LEFT_CM = 1.5
PAYMENT_RIGHT_CM = 1.5


def _add_totals_block(
    doc: DocumentType,
    items_total: Decimal,
    items_count: int,
    installation: Decimal,
    delivery: Decimal,
    unloading: Decimal,
    grand_total: Decimal,
) -> None:
    _add_heading(
        doc,
        "Стоимость заказа:",
        bold=True,
        size=HEADING_FONT_SIZE,
        left_cm=COST_LEFT_CM,
        right_cm=COST_RIGHT_CM,
    )

    # Стоимость заказа — единая таблица (позиции + «Итого:» одной строкой).
    rows = [
        (f"Стоимость изделий, {items_count} шт", _fmt_money(items_total), False),
        ("Монтаж", _fmt_money(installation), False),
        ("Доставка", _fmt_money(delivery), False),
        ("Разгрузка", _fmt_money(unloading), False),
        ("Итого:", _fmt_money(grand_total), True),
    ]

    table = doc.add_table(rows=len(rows), cols=2)
    table.style = "Table Grid"
    total_width = PAGE_WIDTH_CM - COST_LEFT_CM - COST_RIGHT_CM
    label_w = total_width * 0.7
    value_w = total_width - label_w
    _set_table_edge_indent(table, COST_LEFT_CM, COST_RIGHT_CM, [label_w, value_w])

    for i, (label, value, is_total) in enumerate(rows):
        c0, c1 = table.rows[i].cells

        p0 = c0.paragraphs[0]
        _zero_spacing(p0)
        r0 = p0.add_run(label)
        _set_font(r0, bold=is_total)

        p1 = c1.paragraphs[0]
        _zero_spacing(p1)
        r1 = p1.add_run(value)
        _set_font(r1, bold=True)

    spacer = doc.add_paragraph()
    _zero_spacing(spacer)

    _add_heading(
        doc,
        "График платежей:",
        bold=True,
        size=HEADING_FONT_SIZE,
        left_cm=PAYMENT_LEFT_CM,
        right_cm=PAYMENT_RIGHT_CM,
    )

    pay_table = doc.add_table(rows=2, cols=2)
    pay_table.style = "Table Grid"
    pay_total_width = PAGE_WIDTH_CM - PAYMENT_LEFT_CM - PAYMENT_RIGHT_CM
    pay_label_w = pay_total_width * 0.7
    pay_value_w = pay_total_width - pay_label_w
    _set_table_edge_indent(
        pay_table,
        PAYMENT_LEFT_CM,
        PAYMENT_RIGHT_CM,
        [pay_label_w, pay_value_w],
    )

    first_payment = (grand_total * Decimal("0.7")).quantize(Decimal("0.01"))
    second_payment = (grand_total - first_payment).quantize(Decimal("0.01"))
    pay_rows = [
        ("Первый платеж, 70% (предоплата)", _fmt_money(first_payment)),
        ("Второй платеж, 30% (по готовности изделий)", _fmt_money(second_payment)),
    ]
    for i, (label, value) in enumerate(pay_rows):
        c0, c1 = pay_table.rows[i].cells

        p0 = c0.paragraphs[0]
        _zero_spacing(p0)
        r0 = p0.add_run(label)
        _set_font(r0)

        p1 = c1.paragraphs[0]
        _zero_spacing(p1)
        r1 = p1.add_run(value)
        _set_font(r1, bold=True)


# Точка входа
def build_commercial_proposal(
    order,
    *,
    company_name: str = "Индивидуальный предприниматель Дубров Игорь Викторович",
    inn: str = "000000000",
    ogrn: str = "0000000000",
    proposal_date: Optional[str] = None,
    shipment_term: str = "",
) -> io.BytesIO:
    """
    Строит .docx коммерческого предложения по заказу `order` (экземпляр
    orders.models.Order) и возвращает BytesIO с готовым файлом.

    Изделия заказа находятся сами — через related-менеджеры
    order.portal_set (calculate.models.Portal) и order.glukhar_set
    (calculate.models.Glukhar). Каждая строка Portal/Glukhar формирует
    свой блок в документе; поле `amount` этой строки — это количество
    одинаковых изделий данного типа/размера в заказе.
    """
    if proposal_date is None:
        # НАСТРОЙКА: если нужна другая дата (например order.created_at),
        # передайте её явно через параметр proposal_date.
        proposal_date = date.today().strftime("%d.%m.%y")

    doc = Document()
    _configure_normal_style(doc)

    section = doc.sections[0]
    section.page_width = Cm(PAGE_WIDTH_CM)
    section.page_height = Cm(PAGE_HEIGHT_CM)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(1.0)
    # Поля слева/справа = 0: все отступы блоков ниже заданы абсолютно
    # («от края листа»), а не «от границ текста», поэтому поля страницы
    # сами по себе не должны создавать дополнительный отступ.
    section.left_margin = Cm(0)
    section.right_margin = Cm(0)

    _add_document_header(doc, company_name, inn, ogrn, proposal_date, shipment_term)

    portals = list(
        order.portal_set.select_related(
            "scheme",
            "wood_type",
            "hardware_type",
            "color_type",
        ).all(),
    )
    glukhars = list(order.glukhar_set.select_related("wood_type", "color_type").all())

    items_total = Decimal("0")
    items_count = 0

    for portal in portals:
        items_total += get_product_price(portal)
        items_count += portal.amount
        _add_portal_section(doc, portal)

    for i, glukhar in enumerate(glukhars):
        items_total += get_product_price(glukhar)
        items_count += glukhar.amount
        _add_glukhar_section(doc, glukhar, i)

    installation = _to_decimal(order.installation)
    delivery = _to_decimal(order.delivery)
    unloading = _to_decimal(order.unloading)

    grand_total = _to_decimal(order.total_sum) or (
        items_total + installation + delivery + unloading
    )

    _add_totals_block(
        doc,
        items_total=items_total,
        items_count=items_count,
        installation=installation,
        delivery=delivery,
        unloading=unloading,
        grand_total=grand_total,
    )

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer
